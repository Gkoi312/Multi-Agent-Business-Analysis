import os
import sys
import re
import time
from datetime import datetime
from typing import Any
from pathlib import Path
from dotenv import load_dotenv
from langgraph.types import Send

current_dir = os.path.dirname(os.path.abspath(__file__))
project_root = os.path.abspath(os.path.join(current_dir, "../../"))
sys.path.append(project_root)

from langgraph.graph import StateGraph, START, END
from langgraph.checkpoint.memory import MemorySaver
from langchain_core.messages import HumanMessage, SystemMessage
from langchain_community.tools.tavily_search import TavilySearchResults

from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas

from app.schemas.models import (
    AnalystPlan,
    CoverageGoal,
    Perspectives,
    ResearchPlan,
    ReviewFinding,
    ReviewSummary,
    GenerateAnalystsState,
    ResearchGraphState,
)
from app.utils.model_loader import ModelLoader
from app.services.skill_registry import SkillRegistry
from app.workflows.interview_workflow import InterviewGraphBuilder
from app.prompt_lib.prompt_locator import (
    CREATE_ANALYSTS_PROMPT,
    INTRO_CONCLUSION_INSTRUCTIONS,
    REPORT_WRITER_INSTRUCTIONS,
)
from app.logger import GLOBAL_LOGGER
from app.exception.custom_exception import ResearchAnalystException
from app.config import GENERATED_REPORT_DIR


class AutonomousReportGenerator:
    """
    Handles the end-to-end autonomous report generation workflow using LangGraph.
    """

    def __init__(self, llm):
        self.llm = llm
        self.memory = MemorySaver()
        load_dotenv()
        tavily_api_key = os.getenv("TAVILY_API_KEY")
        if not tavily_api_key:
            raise ResearchAnalystException(
                "TAVILY_API_KEY is missing. Please set it in your .env file.",
                ValueError("Missing TAVILY_API_KEY environment variable"),
            )
        self.tavily_search = TavilySearchResults(
            tavily_api_key=tavily_api_key
        )
        self.skill_registry = SkillRegistry(Path(project_root) / "skills")
        self.logger = GLOBAL_LOGGER.bind(module="AutonomousReportGenerator")

    @staticmethod
    def _value(obj: Any, key: str, default=None):
        if isinstance(obj, dict):
            return obj.get(key, default)
        return getattr(obj, key, default)

    @staticmethod
    def _fallback_domain_memory() -> list[dict[str, Any]]:
        return [
            {
                "memory_id": "dd-framework-1",
                "category": "procedure_memory",
                "title": "Due diligence core framework",
                "content": "Cover business model, scale/growth, risk, and final recommendations; tie conclusions to evidence.",
                "tags": ["due_diligence", "framework"],
            },
            {
                "memory_id": "risk-rubric-1",
                "category": "reference_memory",
                "title": "Risk severity rubric",
                "content": "Classify risks High/Medium/Low using impact scope, likelihood, and reversibility.",
                "tags": ["risk", "rubric"],
            },
        ]

    @staticmethod
    def _format_skill_catalog(skills: list[dict[str, Any]]) -> str:
        if not skills:
            return ""
        return "\n\n".join(
            [
                (
                    f"- {AutonomousReportGenerator._value(skill, 'name', '')} "
                    f"({AutonomousReportGenerator._value(skill, 'id', '')})\n"
                    f"  Objective: {AutonomousReportGenerator._value(skill, 'objective', '')}\n"
                    f"  Focus: {', '.join(AutonomousReportGenerator._value(skill, 'focus_areas', []) or [])}"
                )
                for skill in skills
            ]
        )

    def classify_company_type(self, state: ResearchGraphState):
        hint = str(state.get("industry_pack", "") or "").strip().lower()
        supported = set(self.skill_registry.list_industry_packs())
        if hint and hint in supported:
            return {
                "company_type": hint,
                "company_type_confidence": 1.0,
                "company_type_source": "manual",
                "workflow_events": [{"event": "company_type.classified", "payload": {"industry_pack": hint, "source": "manual"}}],
            }
        return {
            "company_type": "unknown",
            "company_type_confidence": 0.0,
            "company_type_source": "fallback",
            "workflow_events": [{"event": "company_type.classified", "payload": {"industry_pack": "unknown", "source": "fallback"}}],
        }

    def assemble_skills(self, state: ResearchGraphState):
        company_type = str(state.get("company_type", "unknown") or "unknown")
        pack = self.skill_registry.load_skill_pack(company_type)
        role_skills = list(pack.get("role_skills", []) or [])
        research_skills = list(pack.get("research_skills", []) or [])
        mappings = list(pack.get("mappings", []) or [])
        source_policies = list(pack.get("source_policies", []) or [])
        source_policy_map = {str(p.get("id", "")): p for p in source_policies if str(p.get("id", ""))}
        selected = role_skills
        domain_memory = list(pack.get("domain_memory", []) or [])
        if not domain_memory:
            domain_memory = self._fallback_domain_memory()
        selected_research_skills: list[dict[str, Any]] = []
        skill_mapping: dict[str, Any] = {}
        if selected and research_skills:
            research_by_id = {str(r.get("id", "")): r for r in research_skills if str(r.get("id", ""))}
            for mapping in mappings:
                role_id = str(mapping.get("role_skill_id", ""))
                research_id = str(mapping.get("research_skill_id", ""))
                if role_id and research_id and research_id in research_by_id:
                    skill_mapping[role_id] = research_id
            selected_ids = {str(s.get("id", "")) for s in selected}
            for role_id in selected_ids:
                rid = str(skill_mapping.get(role_id, ""))
                if rid and rid in research_by_id:
                    selected_research_skills.append(research_by_id[rid])
        return {
            "skill_bundle": selected,
            "research_skills": selected_research_skills,
            "skill_mapping": skill_mapping,
            "source_policy_map": source_policy_map,
            "domain_memory": domain_memory,
            "workflow_events": [
                {
                    "event": "skills.assembled",
                    "payload": {
                        "industry_pack": company_type,
                        "role_skill_count": len(selected),
                        "research_skill_count": len(selected_research_skills),
                    },
                }
            ],
        }

    @staticmethod
    def _extract_usage(message) -> dict:
        usage = {}
        response_meta = getattr(message, "response_metadata", {}) or {}
        usage_meta = getattr(message, "usage_metadata", {}) or {}
        token_usage = response_meta.get("token_usage", {}) if isinstance(response_meta, dict) else {}
        if not isinstance(token_usage, dict):
            token_usage = {}
        usage["prompt_tokens"] = (
            usage_meta.get("input_tokens")
            or token_usage.get("prompt_tokens")
            or token_usage.get("input_tokens")
            or response_meta.get("input_tokens")
            or response_meta.get("prompt_tokens")
            or 0
        )
        usage["completion_tokens"] = (
            usage_meta.get("output_tokens")
            or token_usage.get("completion_tokens")
            or token_usage.get("output_tokens")
            or response_meta.get("output_tokens")
            or response_meta.get("completion_tokens")
            or 0
        )
        usage["total_tokens"] = (
            usage_meta.get("total_tokens")
            or token_usage.get("total_tokens")
            or response_meta.get("total_tokens")
            or usage["prompt_tokens"] + usage["completion_tokens"]
        )
        return usage

    def _resolve_analyst_skill_id(
        self,
        analyst,
        skill_bundle: list[dict[str, Any]],
    ) -> tuple[str, bool]:
        """
        - skill_id in YAML role_skills → keep (use that card).
        - skill_id empty → keep empty (no card); not an error.
        - skill_id non-empty but not in YAML → clear to ""; caller should log (model mistake).
        Returns (skill_id, cleared_invalid); cleared_invalid only True for the last case.
        """
        valid = {
            str(self._value(s, "id", "") or "").strip()
            for s in skill_bundle
            if str(self._value(s, "id", "") or "").strip()
        }
        cand = str(analyst.skill_id or "").strip()
        if cand and cand in valid:
            return cand, False
        return "", bool(cand)

    # ----------------------------------------------------------------------
    def create_analyst(self, state: GenerateAnalystsState):
        """Generate analyst personas based on research brief and feedback."""
        research_query = state["research_query"]
        max_analysts = state["max_analysts"]
        human_analyst_feedback = state.get("human_analyst_feedback", "")
        skill_bundle = state.get("skill_bundle", []) or []
        try:
            self.logger.info("Creating analyst personas", research_query=research_query)
            structured_llm = self.llm.with_structured_output(Perspectives)
            system_prompt = CREATE_ANALYSTS_PROMPT.render(
                research_query=research_query,
                max_analysts=max_analysts,
                skill_count=len(skill_bundle),
                human_analyst_feedback=human_analyst_feedback,
                skill_catalog=self._format_skill_catalog(skill_bundle),
            )
            started_at = time.perf_counter()
            analysts = structured_llm.invoke([
                SystemMessage(content=system_prompt),
                HumanMessage(content="Generate this set of analyst personas."),
            ])
            enriched_analysts = []
            for idx, analyst in enumerate(analysts.analysts):
                skill_id, cleared_invalid = self._resolve_analyst_skill_id(
                    analyst, skill_bundle
                )
                if cleared_invalid:
                    self.logger.warning(
                        "Analyst skill_id invalid; cleared (not in role_skills)",
                        analyst_name=analyst.name,
                        index=idx,
                        model_skill_id=analyst.skill_id,
                    )
                enriched_analysts.append(analyst.model_copy(update={"skill_id": skill_id}))
            latency_ms = int((time.perf_counter() - started_at) * 1000)
            usage = self._extract_usage(analysts)
            self.logger.info("Analysts created", count=len(analysts.analysts))
            return {
                "analysts": enriched_analysts,
                "llm_metrics": [
                    {
                        "node": "report.create_analyst",
                        "latency_ms": latency_ms,
                        "prompt_tokens": usage["prompt_tokens"],
                        "completion_tokens": usage["completion_tokens"],
                        "total_tokens": usage["total_tokens"],
                    }
                ],
            }
        except Exception as e:
            self.logger.error("Error creating analysts", error=str(e))
            raise ResearchAnalystException("Failed to create analysts", e)

    # ----------------------------------------------------------------------
    def human_feedback(self):
        """Pause node for human analyst feedback."""
        try:
            self.logger.info("Awaiting human feedback")
        except Exception as e:
            self.logger.error("Error during feedback stage", error=str(e))
            raise ResearchAnalystException("Human feedback node failed", e)

    # ----------------------------------------------------------------------
    def regenerate_analyst(self, state: GenerateAnalystsState):
        """Re-generate analyst personas after receiving human feedback."""
        return self.create_analyst(state)

    def plan_research(self, state: ResearchGraphState):
        if not bool(state.get("planner_enabled", True)):
            return {
                "research_plan": ResearchPlan(summary="Planner skipped; proceeding to parallel research."),
                "workflow_events": [{"event": "planner.skipped", "payload": {}}],
            }
        analysts = state.get("analysts", []) or []
        research_skills = state.get("research_skills", []) or []
        skill_mapping = state.get("skill_mapping", {}) or {}
        source_policy_map = state.get("source_policy_map", {}) or {}
        plans: list[AnalystPlan] = []
        coverage = [
            CoverageGoal(theme="Business model & competition", why_it_matters="Core value and moat."),
            CoverageGoal(theme="Scale & growth", why_it_matters="Growth trajectory and timing."),
            CoverageGoal(theme="Risk & compliance", why_it_matters="Downside and regulatory exposure."),
        ]
        for analyst in analysts:
            role_skill_id = str(analyst.skill_id or "")
            research_skill_id = str(skill_mapping.get(role_skill_id, "") or "")
            research_skill = next((r for r in research_skills if str(r.get("id", "")) == research_skill_id), {})
            policy_id = str(research_skill.get("source_policy_id", "") or "")
            policy = source_policy_map.get(policy_id, {})
            key_questions = list(research_skill.get("question_templates", []) or [])
            if not key_questions:
                key_questions = [
                    f"From the '{analyst.role}' lens, what verifiable facts or data most support your view (be specific on source type or metrics)?",
                    "Under public information and stated assumptions, which link is most uncertain and would materially change conclusions if wrong?",
                ]
            policy_label = str(policy.get("label", "Default search policy") or "Default search policy")
            plans.append(
                AnalystPlan(
                    analyst_name=analyst.name,
                    skill_id=role_skill_id,
                    research_skill_id=research_skill_id,
                    brief=f"Collect due diligence evidence and produce the memo section from the '{analyst.role}' angle. Suggested policy: {policy_label}",
                    key_questions=key_questions,
                    source_policy=policy,
                )
            )
        plan = ResearchPlan(
            summary="Research plan generated per analyst, covering business, growth, and risk.",
            coverage_goals=coverage,
            analyst_plans=plans,
        )
        return {
            "research_plan": plan,
            "workflow_events": [{"event": "planner.completed", "payload": {"analyst_plans": len(plans)}}],
        }

    # ----------------------------------------------------------------------
    def write_report(self, state: ResearchGraphState):
        """Compile all report sections into unified content."""
        sections = state.get("sections", [])
        research_query = state.get("research_query", "")

        try:
            if not sections:
                sections = ["No sections were generated; check whether the interview stage completed successfully."]
            self.logger.info("Writing report", research_query=research_query)
            system_prompt = REPORT_WRITER_INSTRUCTIONS.render(research_query=research_query)
            started_at = time.perf_counter()
            report = self.llm.invoke([
                SystemMessage(content=system_prompt),
                HumanMessage(content="\n\n".join(sections))
            ])
            latency_ms = int((time.perf_counter() - started_at) * 1000)
            usage = self._extract_usage(report)
            self.logger.info("Report written successfully")
            return {
                "content": report.content,
                "llm_metrics": [
                    {
                        "node": "report.write_report",
                        "latency_ms": latency_ms,
                        "prompt_tokens": usage["prompt_tokens"],
                        "completion_tokens": usage["completion_tokens"],
                        "total_tokens": usage["total_tokens"],
                    }
                ],
            }
        except Exception as e:
            self.logger.error("Error writing main report", error=str(e))
            raise ResearchAnalystException("Failed to write main report", e)

    # ----------------------------------------------------------------------
    def write_introduction(self, state: ResearchGraphState):
        """Generate the report introduction."""
        try:
            sections = state["sections"]
            research_query = state["research_query"]
            formatted_str_sections = "\n\n".join([f"{s}" for s in sections])
            self.logger.info("Generating introduction", research_query=research_query)
            system_prompt = INTRO_CONCLUSION_INSTRUCTIONS.render(
                research_query=research_query, formatted_str_sections=formatted_str_sections
            )
            started_at = time.perf_counter()
            intro = self.llm.invoke([
                SystemMessage(content=system_prompt),
                HumanMessage(content="Write the report introduction.")
            ])
            latency_ms = int((time.perf_counter() - started_at) * 1000)
            usage = self._extract_usage(intro)
            self.logger.info("Introduction generated", length=len(intro.content))
            return {
                "introduction": intro.content,
                "llm_metrics": [
                    {
                        "node": "report.write_introduction",
                        "latency_ms": latency_ms,
                        "prompt_tokens": usage["prompt_tokens"],
                        "completion_tokens": usage["completion_tokens"],
                        "total_tokens": usage["total_tokens"],
                    }
                ],
            }
        except Exception as e:
            self.logger.error("Error generating introduction", error=str(e))
            raise ResearchAnalystException("Failed to generate introduction", e)

    # ----------------------------------------------------------------------
    def write_conclusion(self, state: ResearchGraphState):
        """Generate the conclusion section."""
        try:
            sections = state["sections"]
            research_query = state["research_query"]
            formatted_str_sections = "\n\n".join([f"{s}" for s in sections])
            self.logger.info("Generating conclusion", research_query=research_query)
            system_prompt = INTRO_CONCLUSION_INSTRUCTIONS.render(
                research_query=research_query, formatted_str_sections=formatted_str_sections
            )
            started_at = time.perf_counter()
            conclusion = self.llm.invoke([
                SystemMessage(content=system_prompt),
                HumanMessage(content="Write the report conclusion.")
            ])
            latency_ms = int((time.perf_counter() - started_at) * 1000)
            usage = self._extract_usage(conclusion)
            self.logger.info("Conclusion generated", length=len(conclusion.content))
            return {
                "conclusion": conclusion.content,
                "llm_metrics": [
                    {
                        "node": "report.write_conclusion",
                        "latency_ms": latency_ms,
                        "prompt_tokens": usage["prompt_tokens"],
                        "completion_tokens": usage["completion_tokens"],
                        "total_tokens": usage["total_tokens"],
                    }
                ],
            }
        except Exception as e:
            self.logger.error("Error generating conclusion", error=str(e))
            raise ResearchAnalystException("Failed to generate conclusion", e)

    def review_report(self, state: ResearchGraphState):
        if not bool(state.get("review_enabled", True)):
            return {
                "report_review": ReviewSummary(status="skipped", summary="Full report review skipped."),
                "workflow_events": [{"event": "review.report.skipped", "payload": {}}],
            }
        findings: list[ReviewFinding] = []
        content = state.get("content", "") or ""
        has_risk_level = (
            "风险等级" in content
            or "Risk level" in content
            or re.search(r"Risk\s*level\s*[:：]", content, re.IGNORECASE)
        )
        if not has_risk_level:
            findings.append(
                ReviewFinding(
                    severity="medium",
                    title="Risk level labels missing",
                    detail='Main body does not contain "Risk level" markers.',
                    suggested_fix='In "Risk Assessment", label each risk with Risk level: High / Medium / Low.',
                )
            )
        if "## Sources" not in content and "## 信息来源" not in content:
            findings.append(
                ReviewFinding(
                    severity="high",
                    title="Sources section missing",
                    detail='Main body has no "## Sources" section.',
                    suggested_fix="Add a consolidated source list aligned with in-text [n] citations.",
                )
            )
        if "\n## Sources\n" in content:
            body_main, sources_blob = content.rsplit("\n## Sources\n", 1)
            cite_nums = [int(x) for x in re.findall(r"\[(\d+)\]", body_main)]
            max_n = max(cite_nums) if cite_nums else 0
            source_lines = sum(
                1
                for line in sources_blob.strip().splitlines()
                if re.match(r"^\s*\[\d+\]\s+\S", line.strip())
                or re.match(r"^\s*\[\d+\]$", line.strip())
            )
            if max_n >= 2 and source_lines < max_n:
                findings.append(
                    ReviewFinding(
                        severity="high",
                        title="Sources list shorter than in-text citations",
                        detail=(
                            f'Body cites up to [{max_n}] but "## Sources" appears to have '
                            f"{source_lines} numbered entr{'y' if source_lines == 1 else 'ies'}."
                        ),
                        suggested_fix=(
                            "Under ## Sources, add one [n] line per distinct citation index used in the body "
                            f"(1…{max_n}), with full reference text for each."
                        ),
                    )
                )
        status = "pass" if not findings else "needs_revision"
        summary = "Basic quality check passed." if status == "pass" else "Report needs revision."
        return {
            "report_review": ReviewSummary(status=status, summary=summary, findings=findings),
            "review_notes": [{"scope": "report", "status": status, "findings": len(findings)}],
            "workflow_events": [{"event": "review.report.completed", "payload": {"status": status}}],
        }

    # ----------------------------------------------------------------------
    def finalize_report(self, state: ResearchGraphState):
        """Assemble introduction, content, and conclusion into final report."""
        try:
            content = state["content"]
            self.logger.info("Finalizing report compilation")
            if content.startswith("## Insights"):
                content = content.strip("## Insights")

            # Strip trailing Sources block from main body, then append once at the end (intro → body → conclusion → Sources)
            sources_blob = None
            if "\n## Sources\n" in content:
                content, sources_blob = content.rsplit("\n## Sources\n", 1)

            final_report = (
                state["introduction"] + "\n\n---\n\n" +
                content + "\n\n---\n\n" +
                state["conclusion"]
            )
            if sources_blob:
                final_report += "\n\n## Sources\n" + sources_blob.strip()

            self.logger.info("Report finalized")
            return {"final_report": final_report}
        except Exception as e:
            self.logger.error("Error finalizing report", error=str(e))
            raise ResearchAnalystException("Failed to finalize report", e)

    # ----------------------------------------------------------------------
    def save_report(self, final_report: str, report_name: str,
                    format: str = "docx"):
        """Save the report as DOCX or PDF, each in its own subfolder."""
        try:
            self.logger.info("Saving report", report_name=report_name, format=format)
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            safe_report_name = re.sub(r'[\\/*?:"<>|]', "_", report_name)
            base_name = f"{safe_report_name.replace(' ', '_')}_{timestamp}"

            # Root folder (always inside configured app root)
            root_dir = os.fspath(GENERATED_REPORT_DIR)

            # Create subfolder for this report
            report_folder = os.path.join(root_dir, base_name)
            os.makedirs(report_folder, exist_ok=True)

            # Final file path inside that folder
            file_path = os.path.join(report_folder, f"{base_name}.{format}")

            if format == "docx":
                self._save_as_docx(final_report, file_path)
            elif format == "pdf":
                self._save_as_pdf(final_report, file_path)
            else:
                raise ValueError("Invalid format. Use 'docx' or 'pdf'.")

            self.logger.info("Report saved successfully", path=file_path)
            return file_path

        except Exception as e:
            self.logger.error("Error saving report", error=str(e))
            raise ResearchAnalystException("Failed to save report file", e)

    # ----------------------------------------------------------------------
    def _save_as_docx(self, text: str, file_path: str):
        """Helper: save as DOCX."""
        try:
            doc = Document()
            for line in text.split("\n"):
                if line.startswith("# "):
                    doc.add_heading(line[2:], level=1)
                elif line.startswith("## "):
                    doc.add_heading(line[3:], level=2)
                elif line.startswith("### "):
                    doc.add_heading(line[4:], level=3)
                else:
                    doc.add_paragraph(line)
            doc.save(file_path)
        except Exception as e:
            self.logger.error("DOCX save failed", path=file_path, error=str(e))
            raise ResearchAnalystException("Error saving DOCX report", e)

    def _save_as_pdf(self, text: str, file_path: str):
        """Helper: save as PDF with centered text block, wrapping, and clean layout."""
        from textwrap import wrap
        try:
            c = canvas.Canvas(file_path, pagesize=letter)
            width, height = letter

            # Margins and layout control
            left_margin = 80
            right_margin = 80
            usable_width = width - left_margin - right_margin
            top_margin = 70
            bottom_margin = 60
            y = height - top_margin

            # Fonts and styles
            normal_font = "Helvetica"
            bold_font = "Helvetica-Bold"
            line_height = 15

            # Title centered at top
            lines = text.split("\n")
            for raw_line in lines:
                line = raw_line.strip()
                if not line:
                    y -= line_height
                    continue

                # Detect headings
                if line.startswith("# "):
                    font = bold_font
                    size = 16
                    line = line[2:].strip()
                elif line.startswith("## "):
                    font = bold_font
                    size = 13
                    line = line[3:].strip()
                else:
                    font = normal_font
                    size = 11

                # Wrap text for readable width
                c.setFont(font, size)
                wrapped_lines = wrap(line, width=int(usable_width / (size * 0.55)))

                for wline in wrapped_lines:
                    # 🔹 Auto new page if near bottom
                    if y < bottom_margin:
                        c.showPage()
                        c.setFont(font, size)
                        y = height - top_margin

                    # 🔹 Compute centered X position
                    text_width = c.stringWidth(wline, font, size)
                    x = (width - text_width) / 2  # center horizontally

                    c.drawString(x, y, wline)
                    y -= line_height

            # Optional footer with page number
            for page_num in range(1, c.getPageNumber() + 1):
                c.setFont("Helvetica", 9)
                c.drawCentredString(width / 2, 25, f"Page {page_num}")

            c.save()
            self.logger.info("Centered PDF saved successfully", path=file_path)

        except Exception as e:
            self.logger.error("PDF save failed", path=file_path, error=str(e))
            raise ResearchAnalystException("Error saving PDF report", e)

    # ----------------------------------------------------------------------
    def build_graph(self):
        """Construct the report generation graph."""
        try:
            self.logger.info("Building report generation graph")
            builder = StateGraph(ResearchGraphState)
            interview_graph = InterviewGraphBuilder(self.llm, self.tavily_search).build()

            def initiate_all_interviews(state: ResearchGraphState):
                research_query = state.get("research_query", "Unnamed due diligence task")
                analysts = state.get("analysts", [])
                research_plan = state.get("research_plan")
                skill_bundle = state.get("skill_bundle", []) or []
                domain_memory = state.get("domain_memory", []) or []
                if not analysts:
                    self.logger.warning("No analysts found — skipping interviews")
                    return END
                analyst_plan_map = {}
                analyst_plans = self._value(research_plan, "analyst_plans", []) if research_plan else []
                if analyst_plans:
                    analyst_plan_map = {
                        self._value(p, "analyst_name", ""): p for p in analyst_plans
                    }
                skill_card_by_id = {
                    str(self._value(s, "id", "") or "").strip(): s
                    for s in skill_bundle
                    if str(self._value(s, "id", "") or "").strip()
                }
                return [
                    Send(
                        "conduct_interview",
                        {
                            "analyst": analyst,
                            "skill_card": skill_card_by_id.get(
                                str(analyst.skill_id or "").strip()
                            ),
                            "assigned_plan": analyst_plan_map.get(analyst.name),
                            "domain_memory": domain_memory,
                            "messages": [HumanMessage(content=f"Let's discuss this due diligence task: {research_query}")],
                            "turn_count": 0,
                            "context": [],
                            "retrieved_sources": [],
                            "router_decisions": [],
                            "review_notes": [],
                            "workflow_events": [],
                            "interview": "",
                            "sections": [],
                            "llm_metrics": [],
                        },
                    )
                    for analyst in analysts
                ]

            builder.add_node("classify_company_type", self.classify_company_type)
            builder.add_node("assemble_skills", self.assemble_skills)
            builder.add_node("create_analyst", self.create_analyst)
            builder.add_node("human_feedback", self.human_feedback)
            builder.add_node("regenerate_analyst", self.regenerate_analyst)
            builder.add_node("plan_research", self.plan_research)
            builder.add_node("start_interviews", lambda state: {})
            builder.add_node("conduct_interview", interview_graph)
            builder.add_node("write_report", self.write_report)
            builder.add_node("write_introduction", self.write_introduction)
            builder.add_node("write_conclusion", self.write_conclusion)
            builder.add_node("review_report", self.review_report)
            builder.add_node("finalize_report", self.finalize_report)

            def route_after_feedback(state: ResearchGraphState):
                feedback = (state.get("human_analyst_feedback", "") or "").strip()
                if feedback:
                    return "regenerate_analyst"
                return "plan_research"

            builder.add_edge(START, "classify_company_type")
            builder.add_edge("classify_company_type", "assemble_skills")
            builder.add_edge("assemble_skills", "create_analyst")
            builder.add_edge("create_analyst", "human_feedback")
            builder.add_conditional_edges(
                "human_feedback",
                route_after_feedback,
                ["regenerate_analyst", "plan_research"],
            )
            builder.add_edge("plan_research", "start_interviews")
            builder.add_conditional_edges(
                "start_interviews",
                initiate_all_interviews,
                ["conduct_interview", END]
            )
            builder.add_edge("regenerate_analyst", "human_feedback")
            builder.add_edge("conduct_interview", "write_report")
            builder.add_edge("conduct_interview", "write_introduction")
            builder.add_edge("conduct_interview", "write_conclusion")
            builder.add_edge(["write_report", "write_introduction", "write_conclusion"], "review_report")
            builder.add_edge("review_report", "finalize_report")
            builder.add_edge("finalize_report", END)

            graph = builder.compile(interrupt_before=["human_feedback"], checkpointer=self.memory)
            self.logger.info("Report generation graph built successfully")
            return graph
        except Exception as e:
            self.logger.error("Error building report graph", error=str(e))
            raise ResearchAnalystException("Failed to build report generation graph", e)


# ----------------------------------------------------------------------
if __name__ == "__main__":
    try:
        llm = ModelLoader().load_llm()
        reporter = AutonomousReportGenerator(llm)
        graph = reporter.build_graph()

        research_query = "Conduct due diligence research on OpenAI"
        thread = {"configurable": {"thread_id": "1"}}
        reporter.logger.info("Starting report generation pipeline", research_query=research_query)

        for _ in graph.stream({"research_query": research_query, "max_analysts": 3}, thread, stream_mode="values"):
            pass

        state = graph.get_state(thread)
        feedback = input("\nEnter feedback, or press Enter to continue: ").strip()
        graph.update_state(thread, {"human_analyst_feedback": feedback}, as_node="human_feedback")

        for _ in graph.stream(None, thread, stream_mode="values"):
            pass

        final_state = graph.get_state(thread)
        final_report = final_state.values.get("final_report")

        if final_report:
            reporter.logger.info("Report generated successfully")
            reporter.save_report(final_report, "OpenAI", "docx")
            reporter.save_report(final_report, "OpenAI", "pdf")
        else:
            reporter.logger.error("No report content generated")

    except Exception as e:
        GLOBAL_LOGGER.error("Fatal error in main execution", error=str(e))
        raise ResearchAnalystException("Autonomous report generation pipeline failed", e)
